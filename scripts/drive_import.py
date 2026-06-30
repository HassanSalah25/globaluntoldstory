#!/usr/bin/env python3
"""Download Google Drive folder with Windows-safe paths and extract docx text."""

import json
import re
import sys
from pathlib import Path

import gdown

FOLDER_ID = "1E_S-dJ_Dhuj3k8a3Z8eXVERb0k7RbxFT"
OUTPUT_DIR = Path(__file__).resolve().parent.parent / "storage" / "drive-import"
EXTRACT_DIR = Path(__file__).resolve().parent.parent / "storage" / "drive-extracted"


def sanitize_path(name: str) -> str:
    """Make path segment safe for Windows."""
    name = re.sub(r'[<>:"/\\|?*]', "-", name)
    name = re.sub(r"\s+", " ", name).strip().rstrip(".")
    return name or "unnamed"


def download_folder():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    url = f"https://drive.google.com/drive/folders/{FOLDER_ID}"

    files = gdown.download_folder(url=url, output=str(OUTPUT_DIR), skip_download=True, quiet=True)
    if not files:
        print("No files found", file=sys.stderr)
        return

    for item in files:
        raw_path = item.path if hasattr(item, "path") else item.get("path", "")
        parts = [sanitize_path(p) for p in raw_path.replace("\\", "/").split("/") if p]
        dest = OUTPUT_DIR.joinpath(*parts) if parts else OUTPUT_DIR / "file.docx"
        dest.parent.mkdir(parents=True, exist_ok=True)
        file_id = item.id if hasattr(item, "id") else item.get("id")

        if dest.exists() and dest.stat().st_size > 500:
            print(f"SKIP {dest}")
            continue

        try:
            gdown.download(id=file_id, output=str(dest), quiet=False)
            print(f"OK {dest}")
        except Exception as e:
            print(f"FAIL {dest}: {e}", file=sys.stderr)


def extract_docx():
    try:
        from docx import Document
    except ImportError:
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
        from docx import Document

    EXTRACT_DIR.mkdir(parents=True, exist_ok=True)
    manifest = []

    for docx_path in sorted(OUTPUT_DIR.rglob("*.docx")):
        rel = docx_path.relative_to(OUTPUT_DIR)
        parts = list(rel.parts)
        folder = "/".join(parts[:-1]) if len(parts) > 1 else ""
        lang_file = parts[-1].replace(".docx", "")
        text_parts = []
        try:
            doc = Document(str(docx_path))
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        text_parts.append(" | ".join(cells))
        except Exception as e:
            text_parts = [f"[EXTRACTION ERROR: {e}]"]

        entry = {
            "folder": folder,
            "language_file": lang_file,
            "relative_path": str(rel),
            "text": "\n\n".join(text_parts),
            "char_count": sum(len(t) for t in text_parts),
        }
        manifest.append(entry)

        out_file = EXTRACT_DIR / (str(rel).replace("\\", "/").replace(".docx", ".json"))
        out_file.parent.mkdir(parents=True, exist_ok=True)
        out_file.write_text(json.dumps(entry, ensure_ascii=False, indent=2), encoding="utf-8")

    (EXTRACT_DIR / "manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(f"Extracted {len(manifest)} documents to {EXTRACT_DIR}")


if __name__ == "__main__":
    cmd = sys.argv[1] if len(sys.argv) > 1 else "all"
    if cmd in ("download", "all"):
        download_folder()
    if cmd in ("extract", "all"):
        extract_docx()
